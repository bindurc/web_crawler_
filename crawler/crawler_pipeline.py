import os
import asyncio
import datetime
from typing import Optional,List
from fastapi import APIRouter, HTTPException
from docx import Document
from pydantic import BaseModel, HttpUrl, Field
from dotenv import load_dotenv
from breath_first import BreathFirstCrawl
from depth_first import DepthFirstCrawl


load_dotenv()


router = APIRouter()

asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())


class CrawlRequest(BaseModel):
    url: str
    strategy:str
    method: str
    depth: int = Field(..., ge=0, le=3)

def save_results_to_docx(strategy,method,results: list[dict]) -> str:
    local_storage_path =  os.path.join(os.path.expanduser("~"), "Downloads","crawl_exports")
    os.makedirs(local_storage_path, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%H%M%S") 
    file_path =  os.path.join(local_storage_path, f"crawl_result_{strategy}_{method}_{timestamp}.docx")

    doc = Document()
    doc.add_heading("Crawled Content", level=1)
    

    for i, item in enumerate(results, start=1):
        doc.add_heading(f"{i}. {item['url']}", level=2)
        doc.add_paragraph(item["text"])

    print(f"Saving DOCX to: {file_path}")
    doc.save(file_path)
    return file_path

@router.post("/")
def start_crawling(request: CrawlRequest):
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        url = request.url
        method = request.method
        strategy = request.strategy
        depth = request.depth
        results = []
        if strategy == "breath first":

            if method == "single":

                crawl_single_page_service = BreathFirstCrawl()
                crawl_single_page = crawl_single_page_service.crawl_single_page(url)
                results = loop.run_until_complete(crawl_single_page)
               
            elif method == "recursive":

                best_first_crawl_service = BreathFirstCrawl()
                best_first_crawl = best_first_crawl_service.breath_first_crawl(url,depth)
                results = loop.run_until_complete(best_first_crawl)

                
        elif strategy == "depth first":

            if method == "single":

                crawl_single_page_service = DepthFirstCrawl()
                crawl_single_page = crawl_single_page_service.crawl_single_page(url)
                results = loop.run_until_complete(crawl_single_page)
                
            if method == "recursive":
                #Use the depth-first crawling strategy

                depth_first_crawl_service = DepthFirstCrawl()
                depth_first_crawl = depth_first_crawl_service.depth_first_crawl(url,depth)
                results = loop.run_until_complete(depth_first_crawl)

                
        else:

            raise HTTPException(status_code=400, detail="Invalid strategy")
            
        save_results_to_docx(strategy,method,results)
        return {
            "message": "Crawling completed and data stored successfully.",
            "strategy":strategy,
            "method":method,
            "pages_crawled": len(results or [])
            
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

